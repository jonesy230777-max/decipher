"""Parse Media_Sales_DNA_Audit_Claude_Code_Brief_Updated.docx into JSON.

Emits four canonical source-of-truth files under reference_docs/:
  media_sales_v1_questions.json     - 31 trait Qs + 3 EQ identity Qs
  media_sales_v1_score_map.json     - per-question option->score (1..5)
  media_sales_v1_narratives.json    - 16 strength/action blocks
  media_sales_v1_archetypes.json    - 8 archetypes + top-2 mapping

Per brief §3-5. Brief text in `docs/` and `reference_docs/`.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
BRIEF = REPO / "reference_docs" / "Media_Sales_DNA_Audit_Claude_Code_Brief_Updated.docx"
OUT_DIR = REPO / "reference_docs"

TRAIT_BY_QNUM = {
    **{q: "cognitive_empathy"  for q in range(1, 7)},    # Q1-Q6
    **{q: "eq"                 for q in range(7, 13)},   # Q7-Q12
    **{q: "pressure_composure" for q in range(13, 19)},  # Q13-Q18
    **{q: "narrative_persuasion" for q in range(19, 25)},# Q19-Q24
    **{q: "eq_bonus"           for q in range(25, 32)},  # Q25-Q31 (refines EQ secondary)
}
DIMENSION_ORDER = [
    "cognitive_empathy", "eq", "pressure_composure", "narrative_persuasion"
]
DIMENSION_LABEL = {
    "cognitive_empathy":   "Cognitive Empathy",
    "eq":                  "Emotional Intelligence",
    "pressure_composure":  "Pressure Composure",
    "narrative_persuasion":"Narrative Persuasion",
}
BAND_THRESHOLDS = {  # per brief §4.2
    "elite":      85,
    "performing": 65,
    "practising": 40,
    "developing": 0,
}
EQ_IDENTITIES = ["regulator", "edge_builder", "observer", "namer"]


def _tables_as_rows(doc: Document) -> list[list[list[str]]]:
    return [[[c.text.strip() for c in row.cells] for row in t.rows] for t in doc.tables]


def _is_score_table(t: list[list[str]]) -> bool:
    if not t or len(t) < 2:
        return False
    header = [c.lower() for c in t[0]]
    return header[:3] == ["opt", "answer text", "score"]


def _is_identity_table(t: list[list[str]]) -> bool:
    if not t or len(t) < 2:
        return False
    header = [c.lower() for c in t[0]]
    return header[:3] == ["opt", "answer text", "identity type"]


def _is_narrative_table(t: list[list[str]]) -> bool:
    if len(t) == 2 and len(t[0]) == 2:
        first = (t[0][0] or "").strip().lower()
        return first == "strength"
    return False


def _is_archetype_table(t: list[list[str]]) -> bool:
    if not t or len(t[0]) != 3:
        return False
    h = [c.lower() for c in t[0]]
    return h[:2] == ["top trait 1", "top trait 2"]


def parse() -> dict[str, object]:
    doc = Document(BRIEF)
    tables = _tables_as_rows(doc)

    # Pull question prompts from paragraphs (Q1..Q31).
    qprompts: dict[int, str] = {}
    for p in doc.paragraphs:
        m = re.match(r"^Q(\d+)\.\s*(.+)$", p.text.strip())
        if m and 1 <= int(m.group(1)) <= 31:
            qprompts[int(m.group(1))] = m.group(2).strip()

    # Score tables for Q1..Q31, in document order.
    score_tables = [t for t in tables if _is_score_table(t)]
    if len(score_tables) < 31:
        raise SystemExit(f"expected 31 score tables, got {len(score_tables)}")
    score_map: dict[str, list[dict]] = {}
    for i, t in enumerate(score_tables[:31], start=1):
        options = []
        for row in t[1:]:
            opt_letter, text, score = (row + ["", "", ""])[:3]
            try:
                s = int(score)
            except ValueError:
                continue
            options.append({"letter": opt_letter, "text": text, "score": s})
        score_map[f"Q{i}"] = options

    # EQ identity table.
    identity_tables = [t for t in tables if _is_identity_table(t)]
    eq_identity: dict[str, list[dict]] = {}
    if identity_tables:
        for row in identity_tables[0][1:]:
            opt_letter, text, ident = (row + ["", "", ""])[:3]
            # Text starts with "EQ Identity Q<n>: ..."
            m = re.match(r"EQ Identity Q(\d+):\s*(.+)$", text)
            if not m:
                continue
            qn = int(m.group(1))
            body = m.group(2).strip()
            key = f"EQ{qn}"
            eq_identity.setdefault(key, []).append({
                "letter": opt_letter,
                "text": body,
                "identity": ident.strip().lower().replace("-", "_"),
            })

    # Narrative blocks. Tables come in document order:
    # CE DEV/PRAC/PERF/ELITE, EQ DEV/.../ELITE, PC, NP -> 16 tables.
    narr_tables = [t for t in tables if _is_narrative_table(t)]
    if len(narr_tables) < 16:
        raise SystemExit(f"expected >=16 narrative tables, got {len(narr_tables)}")
    bands_order = ["developing", "practising", "performing", "elite"]
    narratives: dict[str, dict[str, dict]] = {}
    idx = 0
    for dim in DIMENSION_ORDER:
        narratives[dim] = {}
        for band in bands_order:
            t = narr_tables[idx]
            strength = t[0][1].strip()
            action   = t[1][1].strip()
            narratives[dim][band] = {"strength": strength, "action": action}
            idx += 1

    # Archetype mapping table.
    arch_tables = [t for t in tables if _is_archetype_table(t)]
    if not arch_tables:
        raise SystemExit("no archetype table found")
    arch_pairs: list[dict] = []
    for row in arch_tables[0][1:]:
        if len(row) < 3:
            continue
        arch_pairs.append({"top1": row[0].strip(),
                           "top2": row[1].strip(),
                           "archetype": row[2].strip()})

    # Build questions doc with prompt + trait + scoring options.
    questions = []
    for qn in range(1, 32):
        if qn not in qprompts:
            raise SystemExit(f"missing prompt for Q{qn}")
        questions.append({
            "id": f"Q{qn}",
            "qnum": qn,
            "trait": TRAIT_BY_QNUM[qn],
            "prompt": qprompts[qn],
            "options": score_map[f"Q{qn}"],
        })
    for qn in sorted(eq_identity.keys()):
        questions.append({
            "id": qn,
            "trait": "eq_identity",
            "prompt": f"EQ Identity {qn}",
            "options": eq_identity[qn],
        })

    out = {
        "questions":      questions,
        "score_map":      score_map,
        "eq_identity":    eq_identity,
        "narratives":     narratives,
        "archetypes":     arch_pairs,
        "band_thresholds": BAND_THRESHOLDS,
        "dimensions":     DIMENSION_LABEL,
        "trait_by_qnum":  TRAIT_BY_QNUM,
        "eq_identity_codes": EQ_IDENTITIES,
    }

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    (OUT_DIR / "media_sales_v1_questions.json").write_text(json.dumps(questions, indent=2))
    (OUT_DIR / "media_sales_v1_score_map.json").write_text(json.dumps({"score_map": score_map, "eq_identity": eq_identity}, indent=2))
    (OUT_DIR / "media_sales_v1_narratives.json").write_text(json.dumps(narratives, indent=2))
    (OUT_DIR / "media_sales_v1_archetypes.json").write_text(json.dumps(arch_pairs, indent=2))
    (OUT_DIR / "media_sales_v1_full.json").write_text(json.dumps(out, indent=2))

    print(f"Parsed brief OK:")
    print(f"  Questions:    {len(questions)} ({len(score_map)} trait + {len(eq_identity)} identity)")
    print(f"  Score map:    {sum(len(v) for v in score_map.values())} option rows")
    print(f"  Narratives:   {sum(len(b) for b in narratives.values())} blocks across {len(narratives)} traits")
    print(f"  Archetypes:   {len(arch_pairs)}")
    return out


if __name__ == "__main__":
    parse()
